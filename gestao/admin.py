from .views import gerar_certificado_batismo
from babel.dates import format_date
from collections import defaultdict
from datetime import date
from django.contrib import admin
from django.db.models import F
from django.db.models import Sum, F
from django.db.models.functions import ExtractDay, ExtractMonth
from django.http import HttpResponse
from django.http import HttpResponseRedirect, HttpResponse
from django.shortcuts import redirect
from django.shortcuts import render, redirect
from django.template.response import TemplateResponse
from django.urls import path, reverse
from django.urls import reverse, path
from django.utils.dateformat import DateFormat
from django.utils.dateformat import format
from django.utils.html import format_html
from django.utils.html import mark_safe
from django.utils.translation import gettext_lazy as _
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, NamedStyle, PatternFill
from openpyxl.styles import NamedStyle, Font, Alignment
from django.contrib import admin
from .models import Profissao, Habilidades, Autoreslivros, CadastroContas, CadastroSubcontas
from .models import (EntradaFinanceira, SaidaFinanceira, BalanceteMensal, SaldoAnual, BancoFinanceiro, MovimentoFinanceiro,
                     MembroFinanceiro, Agendapastoral, Agendaigreja, Cadastroestudos, Batismoigreja,
                     ContasPagar, Biblioteca, Patrimonio, Fornecedor, EmprestimoLivro, Biblia)

@admin.register(Biblia)
class BibliaAdmin(admin.ModelAdmin):
    list_display = ('nome_livro', 'capitulo', 'versiculo', 'texto_biblico')
    search_fields = ('nome_livro', 'capitulo', 'versiculo', 'texto_biblico')
    list_filter = ('nome_livro', 'capitulo')
    ordering = ('capitulo','versiculo')

    def has_add_permission(self, request):
        return False

@admin.register(Cadastroestudos)
class CadastroestudosAdmin(admin.ModelAdmin):
    list_display = ('titulo', 'data_registro')
    search_fields = ('titulo', 'data_registro')
    ordering = ['titulo']

    def export_to_word(self, request, queryset):
        # Cria um novo documento do Word
        document = Document()

        # Define o cabeçalho do documento
        document.add_heading('Cadastro de Esboços de Mensagens', 0)

        # Adiciona dados
        for cadastro in queryset:
            document.add_heading(cadastro.titulo, level=1)
            document.add_paragraph('Descrição: ', style='Intense Quote')
            document.add_paragraph(cadastro.descricao)
            document.add_paragraph('Data de Registro: ' + cadastro.data_registro.strftime('%d/%m/%Y'), style='Intense Quote')

            # Adiciona um espaço entre os registros
            document.add_paragraph('\n')

        # Prepara a resposta HTTP
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=CadastroEstudos.docx'
        document.save(response)
        return response

    export_to_word.short_description = "Exportar Esboços de Mensagens para Word"
    actions = ['export_to_word']

@admin.register(Agendapastoral)
class AgendapastoralAdmin(admin.ModelAdmin):
    list_display = ('descricao', 'data_evento')
    search_fields = ('descricao', 'data_evento')
    ordering = ['descricao']

    def export_to_excel(self, request, queryset):
        # Cria um novo workbook e uma worksheet ativa
        wb = Workbook()
        ws = wb.active
        ws.title = "Agenda Pastoral"

        # Define o cabeçalho
        ws.append(['Descrição', 'Data do Evento'])

        # Estilos personalizados
        date_style = NamedStyle(name="date_style", number_format='DD/MM/YYYY')
        alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)

        # Adiciona dados
        for agendapastoral in queryset:
            row = [agendapastoral.descricao, agendapastoral.data_evento]
            ws.append(row)
            # Aplica estilos às células
            ws.cell(row=ws.max_row, column=1).alignment = alignment
            ws.cell(row=ws.max_row, column=2).style = date_style

        # Ajusta a largura das colunas conforme necessário
        ws.column_dimensions['A'].width = 50  # Ajuste a largura da coluna de descrição
        ws.column_dimensions['B'].width = 20  # Ajuste a largura da coluna de data do evento

        # Prepara a resposta HTTP
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename=AgendaPastoral.xlsx'
        wb.save(response)
        return response

    export_to_excel.short_description = "Exportar Agenda Pastoral para Excel"
    actions = ['export_to_excel']

@admin.register(Agendaigreja)
class AgendaigrejaAdmin(admin.ModelAdmin):
    list_display = ('descricao', 'data_evento')
    search_fields = ('descricao', 'data_evento')
    ordering = ['descricao']  # Adicione esta linha

    def export_to_excel(self, request, queryset):
        # Cria um novo workbook e uma worksheet ativa
        wb = Workbook()
        ws = wb.active
        ws.title = "Agenda da Igreja"

        # Define o cabeçalho
        ws.append(['Descrição', 'Data do Evento'])

        # Estilos personalizados
        date_style = NamedStyle(name="date_style", number_format='DD/MM/YYYY')
        alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)

        # Adiciona dados
        for agendaigreja in queryset:
            row = [agendaigreja.descricao, agendaigreja.data_evento]
            ws.append(row)
            # Aplica estilos às células
            ws.cell(row=ws.max_row, column=1).alignment = alignment
            ws.cell(row=ws.max_row, column=2).style = date_style

        # Ajusta a largura das colunas conforme necessário
        ws.column_dimensions['A'].width = 50  # Ajuste a largura da coluna de descrição
        ws.column_dimensions['B'].width = 20  # Ajuste a largura da coluna de data do evento

        # Prepara a resposta HTTP
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename=Agenda da Igreja.xlsx'
        wb.save(response)
        return response

    export_to_excel.short_description = "Exportar Agenda da Igreja para Excel"
    actions = ['export_to_excel']


# Definir a ação personalizada
def gerar_certificado_action(modeladmin, request, queryset):
    if queryset.count() == 1:
        batismo = queryset.first()
        return gerar_certificado_batismo(request, batismo.id)
    else:
        messages.error(request, "Por favor, selecione apenas um registro para gerar o certificado.")
        return redirect('..')

gerar_certificado_action.short_description = 'Gerar Certificado de Batismo'

class BatismoigrejaAdmin(admin.ModelAdmin):
    list_display = ('nome', 'formatted_data_batismo')
    search_fields = ('nome', 'data_batismo')
    actions = [gerar_certificado_action, 'export_to_excel']
    ordering = ['data_batismo', 'nome']

    def formatted_data_batismo(self, obj):
        return format(obj.data_batismo, 'd/m/Y')
    formatted_data_batismo.short_description = 'Data de Batismo'

    def export_to_excel(self, request, queryset):
        # Cria um novo workbook e uma worksheet ativa
        wb = Workbook()
        ws = wb.active
        ws.title = "Lista de Batismo"

        # Define o cabeçalho
        ws.append(['Nome', 'Data do Batismo'])

        # Estilos personalizados
        date_style = NamedStyle(name="date_style", number_format='DD/MM/YYYY')
        alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)

        # Adiciona dados
        for batismo in queryset:
            row = [batismo.nome, batismo.data_batismo]
            ws.append(row)
            # Aplica estilos às células
            ws.cell(row=ws.max_row, column=1).alignment = alignment
            ws.cell(row=ws.max_row, column=2).style = date_style

        # Ajusta a largura das colunas conforme necessário
        ws.column_dimensions['A'].width = 50  # Ajuste a largura da coluna de descrição
        ws.column_dimensions['B'].width = 20  # Ajuste a largura da coluna de data do evento

        # Prepara a resposta HTTP
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename=Lista_de_Batismo.xlsx'
        wb.save(response)
        return response

    export_to_excel.short_description = "Exportar Lista de Batizados para Excel"

# Certifique-se de registrar o modelo apenas uma vez
admin.site.register(Batismoigreja, BatismoigrejaAdmin)

# no seu arquivo admin.py

# Admin para Profissao
class ProfissaoAdmin(admin.ModelAdmin):
    list_display = ('nome_prof',)
    search_fields = ('nome_prof',)
    ordering = ['nome_prof']

    def has_module_permission(self, request):
        # Oculta do menu principal
        return False


# Admin para Habilidades
class HabilidadesAdmin(admin.ModelAdmin):
    list_display = ('nome_habilidade',)
    search_fields = ('nome_habilidade',)
    ordering = ['nome_habilidade']

    def has_module_permission(self, request):
        # Oculta do menu principal
        return False


# Registrar os modelos no admin
admin.site.register(Profissao, ProfissaoAdmin)
admin.site.register(Habilidades, HabilidadesAdmin)

class MesAniversarioFilter(admin.SimpleListFilter):
    title = _('Mês de Aniversário')
    parameter_name = 'mes_aniversario'

    def lookups(self, request, model_admin):
        return (
            ('1', _('Janeiro')),
            ('2', _('Fevereiro')),
            ('3', _('Março')),
            ('4', _('Abril')),
            ('5', _('Maio')),
            ('6', _('Junho')),
            ('7', _('Julho')),
            ('8', _('Agosto')),
            ('9', _('Setembro')),
            ('10', _('Outubro')),
            ('11', _('Novembro')),
            ('12', _('Dezembro')),
        )

    def queryset(self, request, queryset):
        if self.value():
            return queryset.filter(data_aniversario__month=int(self.value()))
        return queryset

@admin.register(MembroFinanceiro)
class MembroFinanceiroAdmin(admin.ModelAdmin):
    list_display = ('nome_membro', 'data_aniversario', 'get_cargo', 'get_grupo', 'get_nome_prof', 'get_nome_habilidade')
    search_fields = ('nome_membro', 'data_aniversario', 'cargo', 'grupo', 'nome_prof__nome_prof', 'nome_habilidade__nome_habilidade')
    ordering = ['nome_membro']

    def save_model(self, request, obj, form, change):
        obj.save(request=request)

    def get_cargo(self, obj):
        return obj.cargo
    get_cargo.short_description = 'Cargo'

    def get_grupo(self, obj):
        return obj.grupo
    get_grupo.short_description = 'Grupo'

    def get_nome_prof(self, obj):
        return obj.nome_prof
    get_nome_prof.short_description = 'Profissão'

    def get_nome_habilidade(self, obj):
        return obj.nome_habilidade
    get_nome_habilidade.short_description = 'Habilidade'

    def get_urls(self):
        urls = super().get_urls()
        custom_urls = [
            path('aniversariantes/', self.admin_site.admin_view(self.aniversariantes_view), name='gestao_membrofinanceiro_aniversariantes')
        ]
        return custom_urls + urls

    def aniversariantes_view(self, request):
        mes_atual = date.today().month
        aniversariantes = MembroFinanceiro.objects.filter(data_aniversario__month=mes_atual).annotate(
            dia=F('data_aniversario__day'),
            mes=F('data_aniversario__month')
        ).order_by('mes', 'dia')
        context = dict(
            self.admin_site.each_context(request),
            aniversariantes=aniversariantes,
        )
        return TemplateResponse(request, "admin/aniversariantes.html", context)

    def export_to_excel(self, request, queryset):
        membros = queryset.all()
        wb = Workbook()
        ws1 = wb.active
        ws1.title = 'Por Data de Aniversário'
        membros_ordenados_por_aniversario = membros.annotate(
            day=ExtractDay('data_aniversario'),
            month=ExtractMonth('data_aniversario')
        ).order_by('month', 'day')
        self.add_data_to_worksheet(ws1, membros_ordenados_por_aniversario)
        ws2 = wb.create_sheet(title='Por Nome')
        self.add_data_to_worksheet(ws2, membros.order_by('nome_membro'))
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename=Lista de Membros.xlsx'
        wb.save(response)
        return response

    def add_data_to_worksheet(self, worksheet, members):
        header_font = Font(bold=True)
        header_alignment = Alignment(horizontal='center')
        headers = ['Nome', 'Cargo', 'Número do Telefone', 'Data de Aniversário']
        worksheet.append(headers)
        for cell in worksheet[1]:
            cell.font = header_font
            cell.alignment = header_alignment
        for membro in members:
            data_aniversario = membro.data_aniversario.strftime('%d/%m/%Y') if membro.data_aniversario else ''
            worksheet.append([
                membro.nome_membro,
                membro.cargo,
                membro.numero_telefone,
                data_aniversario
            ])
        worksheet.column_dimensions['A'].width = 50
        worksheet.column_dimensions['B'].width = 20
        worksheet.column_dimensions['C'].width = 20
        worksheet.column_dimensions['D'].width = 20

    export_to_excel.short_description = "Exportar 2 Listas: Data Aniversario e Nome para Excel"
    actions = [export_to_excel]

    def ver_aniversariantes(self, request, queryset):
        return redirect('admin:gestao_membrofinanceiro_aniversariantes')
    ver_aniversariantes.short_description = "Ver Aniversariantes do Mês"

    class Media:
        css = {
            'all': ('css/admin_custom.css',),
        }

class AutoreslivrosAdmin(admin.ModelAdmin):
    list_display = ('nome_autor',)
    search_fields = ('nome_autor',)
    ordering = ['nome_autor']

    def has_module_permission(self, request):
        # Oculta do menu principal
        return False


# Registrar o modelo no admin fora da definição da classe
admin.site.register(Autoreslivros, AutoreslivrosAdmin)

@admin.register(Biblioteca)
class BibliotecaAdmin(admin.ModelAdmin):
    list_display = ('titulo_livro', 'codigo', 'qtdes_livros', 'categoria', 'local', 'get_nome_autor')
    search_fields = ('titulo_livro', 'codigo', 'qtdes_livros', 'categoria', 'local', 'nome_autor__nome_autor')
    ordering = ['titulo_livro']

    def get_nome_autor(self, obj):
        return obj.nome_autor
    get_nome_autor.short_description = 'Nome Autor'

    def export_to_excel(self, request, queryset):
        wb = Workbook()
        ws = wb.active
        ws.append(['Título do Livro', 'Código do Livro', 'Nome do Autor', 'Quantidade de Livros', 'Categoria'])

        # Estilos para o cabeçalho
        header_font = Font(bold=True)
        header_alignment = Alignment(horizontal='center')
        for cell in ws[1]:
            cell.font = header_font
            cell.alignment = header_alignment

        for biblioteca in queryset:
            row = [
                biblioteca.titulo_livro,
                biblioteca.codigo,
                biblioteca.nome_autor,
                biblioteca.qtdes_livros,
                biblioteca.categoria
            ]
            ws.append(row)

        # Ajustar a largura das colunas
        ws.column_dimensions['A'].width = 60
        ws.column_dimensions['B'].width = 20
        ws.column_dimensions['C'].width = 40
        ws.column_dimensions['D'].width = 20
        ws.column_dimensions['E'].width = 30

        # Prepara a resposta HTTP
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = 'attachment; filename=Lista_de_Livros.xlsx'
        wb.save(response)
        return response

    export_to_excel.short_description = "Exportar Listagem de Livros para Excel"
    actions = ['export_to_excel']

@admin.register(EmprestimoLivro)
class EmprestimoLivroAdmin(admin.ModelAdmin):
    list_display = ['get_titulo_livro', 'get_nome_membro', 'get_qtdes_livros', 'formatted_data_emprestimo', 'formatted_data_devolucao']
    search_fields = ('titulo_livro__titulo_livro', 'nome_membro__nome_membro', 'titulo_livro__qtdes_livros', 'data_emprestimo', 'data_devolucao')
    ordering = ['data_emprestimo']

    def formatted_data_emprestimo(self, obj):
        return obj.data_emprestimo.strftime('%d/%m/%Y %H:%M') if obj.data_emprestimo else '-'
    formatted_data_emprestimo.short_description = 'Data do Empréstimo'

    def formatted_data_devolucao(self, obj):
        return obj.data_devolucao.strftime('%d/%m/%Y %H:%M') if obj.data_devolucao else '-'
    formatted_data_devolucao.short_description = 'Data de Devolução'

    def get_titulo_livro(self, obj):
        return obj.titulo_livro.titulo_livro if obj.titulo_livro else '-'
    get_titulo_livro.short_description = 'Título do Livro'

    def get_qtdes_livros(self, obj):
        return obj.titulo_livro.qtdes_livros if obj.titulo_livro else '-'
    get_qtdes_livros.short_description = 'Qtdes de Livros'

    def get_nome_membro(self, obj):
        return obj.nome_membro.nome_membro if obj.nome_membro else '-'
    get_nome_membro.short_description = 'Nome do Membro'

@admin.register(Patrimonio)
class PatrimonioAdmin(admin.ModelAdmin):
    list_display = ('nome_patr', 'codigo_patr', 'local_patr')
    search_fields = ('nome_patr', 'codigo_patr', 'local_patr')
    ordering = ['nome_patr']

    def export_to_excel(self, request, queryset):
        wb = Workbook()
        ws = wb.active
        ws.append(['Nome do Patrimônio', 'Código do Patrimônio', 'Local do Patrimônio'])

        # Estilos para o cabeçalho
        header_font = Font(bold=True)
        header_alignment = Alignment(horizontal='center')
        for cell in ws[1]:
            cell.font = header_font
            cell.alignment = header_alignment

        for patrimonio in queryset:
            row = [patrimonio.nome_patr, patrimonio.codigo_patr, patrimonio.local_patr]
            ws.append(row)

        # Ajustar a largura das colunas
        ws.column_dimensions['A'].width = 60
        ws.column_dimensions['B'].width = 40
        ws.column_dimensions['C'].width = 40

        # Prepara a resposta HTTP
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = 'attachment; filename=Lista_de_Patrimonio.xlsx'
        wb.save(response)
        return response

    export_to_excel.short_description = "Exportar Listagem de Patrimônio para Excel"
    actions = ['export_to_excel']

@admin.register(Fornecedor)
class FornecedorAdmin(admin.ModelAdmin):
    list_display = ('nome_fornec','site','telefone')
    search_fields = ('nome_fornec','site','telefone')
    ordering = ['nome_fornec']  # Adicione esta linha

# Admin para CadastroContas
class CadastroContasAdmin(admin.ModelAdmin):
    list_display = ('nome_conta', 'nr_contabil')
    search_fields = ('nome_conta', 'nr_contabil')
    ordering = ['nome_conta']

    def has_module_permission(self, request):
        # Oculta do menu principal
        return False

    def export_to_excel(self, request, queryset):
        # Cria um novo workbook e uma worksheet ativa
        wb = Workbook()
        ws = wb.active
        ws.title = "Cadastro de Contas"

        # Define o cabeçalho
        ws.append(['Nome da Conta', 'Número da Conta'])

        # Estilos personalizados
        date_style = NamedStyle(name="date_style", number_format='DD/MM/YYYY')
        alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)

        # Adiciona dados
        for cadastrocontas in queryset:
            row = [cadastrocontas.nome_conta, cadastrocontas.nr_contabil]
            ws.append(row)
            # Aplica estilos às células
            ws.cell(row=ws.max_row, column=1).alignment = alignment
            ws.cell(row=ws.max_row, column=2).style = date_style

        # Ajusta a largura das colunas conforme necessário
        ws.column_dimensions['A'].width = 50  # Ajuste a largura da coluna de descrição
        ws.column_dimensions['B'].width = 20  # Ajuste a largura da coluna de data do evento

        # Prepara a resposta HTTP
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename=Listagem_de_Contas.xlsx'
        wb.save(response)
        return response

    export_to_excel.short_description = "Exportar Listagem de Contas para Excel"
    actions = ['export_to_excel']

# Admin para CadastroSubcontas
class CadastroSubcontasAdmin(admin.ModelAdmin):
    list_display = ('nome_subconta', 'nr_subconta')
    search_fields = ('nome_subconta', 'nr_subconta')
    ordering = ['nome_subconta']

    def has_module_permission(self, request):
        # Oculta do menu principal
        return False

    def export_to_excel(self, request, queryset):
        # Cria um novo workbook e uma worksheet ativa
        wb = Workbook()
        ws = wb.active
        ws.title = "Cadastro de Subcontas"

        # Define o cabeçalho
        ws.append(['Nome da Subconta', 'Número da Subconta'])

        # Estilos personalizados
        date_style = NamedStyle(name="date_style", number_format='DD/MM/YYYY')
        alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)

        # Adiciona dados
        for cadastrosubcontas in queryset:
            row = [cadastrosubcontas.nome_subconta, cadastrosubcontas.nr_subconta]
            ws.append(row)
            # Aplica estilos às células
            ws.cell(row=ws.max_row, column=1).alignment = alignment
            ws.cell(row=ws.max_row, column=2).style = date_style

        # Ajusta a largura das colunas conforme necessário
        ws.column_dimensions['A'].width = 50  # Ajuste a largura da coluna de descrição
        ws.column_dimensions['B'].width = 20  # Ajuste a largura da coluna de data do evento

        # Prepara a resposta HTTP
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename=Listagem_de_Subcontas.xlsx'
        wb.save(response)
        return response

    export_to_excel.short_description = "Exportar Listagem de Subcontas para Excel"
    actions = ['export_to_excel']

# Registrar os modelos no admin
admin.site.register(CadastroContas, CadastroContasAdmin)
admin.site.register(CadastroSubcontas, CadastroSubcontasAdmin)
@admin.register(ContasPagar)
class ContasPagarAdmin(admin.ModelAdmin):
    list_display = ('nome_empresa', 'valor', 'descricao', 'data_vencimento', 'data_pagto')
    search_fields = ('nome_empresa', 'valor', 'descricao', 'data_vencimento', 'data_pagto')
    ordering = ['nome_empresa']


@admin.register(BancoFinanceiro)
class BancoFinanceiroAdmin(admin.ModelAdmin):
    list_display = ('nome_banco', 'numero_agencia', 'numero_conta')
    search_fields = ('nome_banco', 'numero_agencia', 'numero_conta')
    ordering = ['nome_banco']


@admin.register(MovimentoFinanceiro)
class MovimentoFinanceiroAdmin(admin.ModelAdmin):
    list_display = ('banco', 'tipo_conta', 'valor')
    search_fields = ('banco', 'tipo_conta', 'valor')
    ordering = ['banco']

    def export_to_excel(self, request, queryset):
        wb = Workbook()
        ws = wb.active
        ws.append(['Banco', 'Tipo de Conta', 'Valor'])

        # Aplicar estilos aos cabeçalhos
        font_bold = Font(bold=True)
        alignment_center = Alignment(horizontal='center')
        border_thin = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))

        for cell in ws[1]:
            cell.font = font_bold
            cell.alignment = alignment_center
            cell.border = border_thin

        for movimento in queryset:
            row = [
                str(movimento.banco) if movimento.banco else '',
                str(movimento.tipo_conta) if movimento.tipo_conta else '',
                movimento.valor
            ]
            ws.append(row)

        # Aplicar estilos aos dados
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
            for cell in row:
                cell.alignment = alignment_center
                cell.border = border_thin

        # Ajustar largura das colunas
        for col in ws.columns:
            max_length = 0
            column = col[0].column_letter  # Obter a letra da coluna
            for cell in col:
                try:
                    if cell.value and len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2)
            ws.column_dimensions[column].width = adjusted_width

        # Prepara a resposta HTTP
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename=Movimentos_Financeiros.xlsx'
        wb.save(response)
        return response

    export_to_excel.short_description = "Exportar Movimentos Financeiros para Excel"
    actions = ['export_to_excel']


# Registre os estilos apenas uma vez
currency_style = NamedStyle(name="currency_style", number_format='R$ #,##0.00')
date_style = NamedStyle(name="date_style", number_format='DD/MM/YYYY')

@admin.register(EntradaFinanceira)
class EntradaFinanceiraAdmin(admin.ModelAdmin):
    list_display = ('nome_conta', 'nome_subconta', 'banco', 'descricao', 'valor', 'formatted_data_entrada')
    search_fields = ('nome_conta__nome_conta', 'nome_subconta__nome_subconta', 'descricao', 'valor')
    ordering = ['data_entrada']


    def formatted_data_entrada(self, obj):
        return obj.data_entrada.strftime('%d/%m/%Y') if obj.data_entrada else ''

    formatted_data_entrada.short_description = 'Data da Entrada'

    def export_to_excel(self, request, queryset):
        wb = Workbook()
        ws = wb.active
        ws.append(['Nome da Conta', 'Nome da Subconta', 'Nome do Banco', 'Descrição', 'Valor', 'Data da Entrada'])

        # Estilos para o cabeçalho
        header_font = Font(bold=True)
        header_alignment = Alignment(horizontal='center')
        for cell in ws[1]:
            cell.font = header_font
            cell.alignment = header_alignment

        for entrada in queryset:
            data_entrada = entrada.data_entrada.strftime('%d/%m/%Y') if entrada.data_entrada else ''
            row = [
                str(entrada.nome_conta) if entrada.nome_conta else '',
                str(entrada.nome_subconta) if entrada.nome_subconta else '',
                str(entrada.banco) if entrada.banco else '',
                entrada.descricao,
                entrada.valor,
                data_entrada
            ]
            ws.append(row)
            # Aplicar os estilos às células corretas
            ws.cell(row=ws.max_row, column=5).style = currency_style
            ws.cell(row=ws.max_row, column=6).style = date_style

        # Ajustar a largura das colunas
        ws.column_dimensions['A'].width = 60
        ws.column_dimensions['B'].width = 40
        ws.column_dimensions['C'].width = 40
        ws.column_dimensions['D'].width = 30
        ws.column_dimensions['E'].width = 20
        ws.column_dimensions['F'].width = 20

        # Prepara a resposta HTTP
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename=Entradas_de_Valores.xlsx'
        wb.save(response)
        return response

    export_to_excel.short_description = "Exportar Entrada de Valores Recebidos para Excel"
    actions = ['export_to_excel']


@admin.register(SaidaFinanceira)
class SaidaFinanceiraAdmin(admin.ModelAdmin):
    list_display = ('nome_conta', 'nome_subconta', 'banco', 'descricao', 'valor', 'formatted_data_saida')
    search_fields = ('nome_conta__nome_conta', 'nome_subconta__nome_subconta', 'descricao', 'valor')
    ordering = ['data_saida']


    def formatted_data_saida(self, obj):
        return obj.data_saida.strftime('%d/%m/%Y') if obj.data_saida else ''

    formatted_data_saida.short_description = 'Data de Pagamento'

    def export_to_excel(self, request, queryset):
        wb = Workbook()
        ws = wb.active
        ws.append(['Nome da Conta', 'Nome da Subconta', 'Nome do Banco', 'Descrição', 'Valor', 'Data do Pagamento'])

        # Estilos para o cabeçalho
        header_font = Font(bold=True)
        header_alignment = Alignment(horizontal='center')
        for cell in ws[1]:
            cell.font = header_font
            cell.alignment = header_alignment

        for saida in queryset:
            data_saida = saida.data_saida.strftime('%d/%m/%Y') if saida.data_saida else ''
            row = [
                str(saida.nome_conta) if saida.nome_conta else '',
                str(saida.nome_subconta) if saida.nome_subconta else '',
                str(saida.banco) if saida.banco else '',
                saida.descricao,
                saida.valor,
                data_saida
            ]
            ws.append(row)
            # Aplicar os estilos às células corretas
            ws.cell(row=ws.max_row, column=5).style = currency_style
            ws.cell(row=ws.max_row, column=6).style = date_style

        # Ajustar a largura das colunas
        ws.column_dimensions['A'].width = 60
        ws.column_dimensions['B'].width = 40
        ws.column_dimensions['C'].width = 40
        ws.column_dimensions['D'].width = 30
        ws.column_dimensions['E'].width = 20
        ws.column_dimensions['F'].width = 20

        # Prepara a resposta HTTP
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename=Pagamentos Realizados de Valores.xlsx'
        wb.save(response)
        return response

    export_to_excel.short_description = "Exportar Pagamentos de Valores Realizados para Excel"
    actions = ['export_to_excel']

@admin.register(BalanceteMensal)
class BalanceteMensalAdmin(admin.ModelAdmin):
    list_display = ('nome_mes', 'ano', 'total_entradas', 'total_saidas', 'saldo_mensal')
    search_fields = ('mes', 'ano', 'total_entradas', 'total_saidas', 'saldo_mensal')
    ordering = ['mes']

    def has_add_permission(self, request):
        return False

    def nome_mes(self, obj):
        nome_meses = {
            1: 'Janeiro',
            2: 'Fevereiro',
            3: 'Março',
            4: 'Abril',
            5: 'Maio',
            6: 'Junho',
            7: 'Julho',
            8: 'Agosto',
            9: 'Setembro',
            10: 'Outubro',
            11: 'Novembro',
            12: 'Dezembro',
        }
        # Check if obj is a dictionary and has the key 'mes'
        mes = obj['mes'] if isinstance(obj, dict) else obj.mes
        return nome_meses.get(mes, 'Mês Desconhecido')

    nome_mes.short_description = 'Mês'

    def calcular_balancete_por_conta(self, queryset):
        ano = queryset[0].ano
        entradas = EntradaFinanceira.objects.filter(
            data_entrada__year=ano
        ).annotate(mes=F('data_entrada__month')).values('mes', 'nome_conta').annotate(total_entradas=Sum('valor'))

        saidas = SaidaFinanceira.objects.filter(
            data_saida__year=ano
        ).annotate(mes=F('data_saida__month')).values('mes', 'nome_conta').annotate(total_saidas=Sum('valor'))

        resultado = defaultdict(lambda: defaultdict(lambda: {'entradas': 0, 'saidas': 0, 'saldo': 0}))

        for entrada in entradas:
            mes = entrada['mes']
            conta = entrada['nome_conta']
            resultado[conta][mes]['entradas'] += entrada['total_entradas']
            resultado[conta][mes]['saldo'] += entrada['total_entradas']

        for saida in saidas:
            mes = saida['mes']
            conta = saida['nome_conta']
            resultado[conta][mes]['saidas'] += saida['total_saidas']
            resultado[conta][mes]['saldo'] -= saida['total_saidas']

        return resultado

    def export_to_excel(self, request, queryset):
        wb = Workbook()
        ws = wb.active
        ws.title = "Balancete Mensal"
        ws.append(['Mês', 'Ano', 'Total de Entradas', 'Total de Saídas', 'Saldo Mensal'])

        # Definir estilo monetário
        currency_style = NamedStyle(name="currency_style", number_format='R$ #,##0.00')

        for balancete in queryset:
            row = [
                self.nome_mes(balancete),
                balancete.ano,
                balancete.total_entradas,
                balancete.total_saidas,
                balancete.saldo_mensal
            ]
            ws.append(row)

            # Aplicar o estilo monetário
            ws.cell(row=ws.max_row, column=3).style = currency_style
            ws.cell(row=ws.max_row, column=4).style = currency_style
            ws.cell(row=ws.max_row, column=5).style = currency_style

        # Ajustar a largura das colunas
        format_excel(ws)

        # Exportar por nr_contabilidade
        resultado_por_conta = self.calcular_balancete_por_conta(queryset)
        ws = wb.create_sheet(title="Por Conta")
        ws.append(['Conta', 'Mês', 'Total de Entradas', 'Total de Saídas', 'Saldo'])

        for conta, meses in resultado_por_conta.items():
            for mes, valores in meses.items():
                ws.append([
                    conta,
                    self.nome_mes({'mes': mes}),  # Passar um dicionário com a chave 'mes'
                    valores['entradas'],
                    valores['saidas'],
                    valores['saldo']
                ])
                ws.cell(row=ws.max_row, column=3).style = currency_style
                ws.cell(row=ws.max_row, column=4).style = currency_style
                ws.cell(row=ws.max_row, column=5).style = currency_style

        # Ajustar a largura das colunas
        format_excel(ws)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename=Balancete_Mensal.xlsx'
        wb.save(response)
        return response

    export_to_excel.short_description = "Exportar Balancete Mensal para Excel"
    actions = ['export_to_excel']

    def format_excel(ws):
        for column in ws.columns:
            max_length = 0
            column = [cell for cell in column]
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except Exception:
                    pass
            adjusted_width = (max_length + 2)
            ws.column_dimensions[column[0].column_letter].width = adjusted_width


@admin.register(SaldoAnual)
class SaldoAnualAdmin(admin.ModelAdmin):
    list_display = ('ano', 'total_entradas_ano', 'total_saidas_ano', 'saldo_anual')
    search_fields = ('ano', 'total_entradas_ano', 'total_saidas_ano', 'saldo_anual')
    ordering = ['ano']  # Adicione esta linha

    def has_add_permission(self, request):
        return False

    def calcular_saldo_anual_por_conta(self, queryset):
        entradas = EntradaFinanceira.objects.filter(
            data_entrada__year=queryset[0].ano
        ).values('nome_conta').annotate(total_entradas=Sum('valor'))

        saidas = SaidaFinanceira.objects.filter(
            data_saida__year=queryset[0].ano
        ).values('nome_conta').annotate(total_saidas=Sum('valor'))

        resultado = {}
        for entrada in entradas:
            conta = entrada['nome_conta']
            if conta not in resultado:
                resultado[conta] = {
                    'entradas': entrada['total_entradas'],
                    'saidas': 0,
                    'saldo': entrada['total_entradas']
                }
            else:
                resultado[conta]['entradas'] += entrada['total_entradas']
                resultado[conta]['saldo'] += entrada['total_entradas']

        for saida in saidas:
            conta = saida['nome_conta']
            if conta not in resultado:
                resultado[conta] = {
                    'entradas': 0,
                    'saidas': saida['total_saidas'],
                    'saldo': -saida['total_saidas']
                }
            else:
                resultado[conta]['saidas'] += saida['total_saidas']
                resultado[conta]['saldo'] -= saida['total_saidas']

        return resultado

    def export_to_excel(self, request, queryset):
        wb = Workbook()
        ws = wb.active
        ws.append(['Ano', 'Total de Entradas Anual', 'Total de Saídas Anual', 'Saldo Anual'])

        # Definir estilo para moeda
        currency_style = NamedStyle(name="currency_style", number_format='R$ #,##0.00')

        for saldo_anual in queryset:
            row = [
                saldo_anual.ano,
                saldo_anual.total_entradas_ano,
                saldo_anual.total_saidas_ano,
                saldo_anual.saldo_anual
            ]
            ws.append(row)

            # Aplicar o estilo monetário
            for cell in ws.iter_rows(min_row=ws.max_row, max_row=ws.max_row, min_col=2, max_col=4):
                for c in cell:
                    c.style = currency_style  # Corrigido aqui

        # Exportar por nr_contabilidade
        resultado_por_conta = self.calcular_saldo_anual_por_conta(queryset)
        ws = wb.create_sheet(title="Por Conta")
        ws.append(['Conta', 'Total de Entradas Anual', 'Total de Saídas Anual', 'Saldo Anual'])

        for conta, valores in resultado_por_conta.items():
            ws.append([
                conta,
                valores['entradas'],
                valores['saidas'],
                valores['saldo']
            ])

            # Aplicar o estilo monetário
            for cell in ws.iter_rows(min_row=ws.max_row, max_row=ws.max_row, min_col=2, max_col=4):
                for c in cell:
                    c.style = currency_style  # Corrigido aqui

        # Ajustar largura das colunas
        for column in ws.columns:
            max_length = 0
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except Exception:
                    pass
            adjusted_width = (max_length + 2)
            ws.column_dimensions[column[0].column_letter].width = adjusted_width

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename=Saldo_Anual.xlsx'
        wb.save(response)
        return response

    export_to_excel.short_description = "Exportar Saldo Anual para Excel"
    actions = ['export_to_excel']

def format_excel(ws):
    for column in ws.columns:
        max_length = 0
        column = [cell for cell in column]
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except Exception:
                pass
        adjusted_width = (max_length + 2)
        ws.column_dimensions[column[0].column_letter].width = adjusted_width
